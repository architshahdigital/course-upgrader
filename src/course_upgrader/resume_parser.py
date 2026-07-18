"""Resume ingestion: extracts raw text from resume files and builds a
structured CandidateProfile via LLM-assisted extraction."""
from __future__ import annotations

from pathlib import Path

from course_upgrader.llm.base import LLMProvider
from course_upgrader.models import CandidateProfile
from course_upgrader.utils import clean_whitespace, extract_json, truncate

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}

EXTRACTION_SYSTEM_PROMPT = """You are an expert technical recruiter and resume parser.
Given the raw text of a candidate's resume, extract a structured JSON profile.
Only output valid JSON, no commentary, no markdown fences.

The JSON schema is:
{
  "name": string | null,
  "skills": string[],       // soft + hard skills, technologies, languages, methodologies
  "tools": string[],        // specific software, libraries, frameworks, platforms
  "job_titles": string[],   // past and current job titles
  "experience_years": number | null,   // total years of professional experience, estimated
  "education": string[]     // degrees, certifications, institutions
}

Be thorough: capture every technology, tool, framework, programming language, and
methodology explicitly mentioned or clearly implied by the resume content."""


def read_resume_text(file_path: str | Path) -> str:
    """Extract raw text from a .pdf, .docx, .txt, or .md resume file."""
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"Resume file not found: {path}")

    suffix = path.suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"Unsupported resume format '{suffix}'. Supported: {sorted(SUPPORTED_EXTENSIONS)}")

    if suffix == ".pdf":
        return _read_pdf(path)
    if suffix == ".docx":
        return _read_docx(path)
    return path.read_text(encoding="utf-8", errors="ignore")


def _read_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RuntimeError(
            "The 'pypdf' package is required to parse PDF resumes. Install it with: pip install pypdf"
        ) from exc

    reader = PdfReader(str(path))
    pages = [page.extract_text() or "" for page in reader.pages]
    text = "\n".join(pages)
    if not text.strip():
        raise ValueError(
            "No extractable text found in PDF. It may be a scanned image; "
            "try exporting a text-based PDF or a .docx/.txt version instead."
        )
    return text


def _read_docx(path: Path) -> str:
    try:
        import docx
    except ImportError as exc:
        raise RuntimeError(
            "The 'python-docx' package is required to parse DOCX resumes. "
            "Install it with: pip install python-docx"
        ) from exc

    document = docx.Document(str(path))
    paragraphs = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            paragraphs.extend(cell.text for cell in row.cells)
    return "\n".join(paragraphs)


def extract_profile(
    raw_text: str,
    llm: LLMProvider,
    career_goal: str | None = None,
) -> CandidateProfile:
    """Use the configured LLM to turn raw resume text into a CandidateProfile."""
    cleaned = clean_whitespace(raw_text)
    if not cleaned:
        raise ValueError("Resume text is empty after cleaning; nothing to extract.")

    response = llm.complete(
        system_prompt=EXTRACTION_SYSTEM_PROMPT,
        user_prompt=truncate(cleaned, max_chars=12000),
        max_tokens=4096,
    )
    data = extract_json(response)

    return CandidateProfile(
        name=data.get("name"),
        raw_text=cleaned,
        skills=list(dict.fromkeys(data.get("skills") or [])),
        tools=list(dict.fromkeys(data.get("tools") or [])),
        job_titles=list(dict.fromkeys(data.get("job_titles") or [])),
        experience_years=data.get("experience_years"),
        education=list(dict.fromkeys(data.get("education") or [])),
        career_goal=career_goal,
    )
